import os
import re
import time
from functools import reduce, wraps

from django.http import HttpResponse
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


import time
def execution_speed(func):
    @wraps(func)
    def wrapper(request, *args, **kwargs):
        start = time.time()
        func(request)
        end = time.time()
        print(f'execution_speed is {end-start}')
        return HttpResponse()
    return wrapper
    
import subprocess

@execution_speed
def python_docs(request): 
    # # Create charts
    
    LEARNER = 'dianaGera'
    
    CHART_1 = 'chart1SettingsGeneral.json'
    
    
    # Replace values in Render Settings for Chart 4 and create new settings file
    cohort_avarage_data = re.compile(r'\[A-Z_A-Z\]')
    learner_score_data = re.compile(r'\[YOUR_SCORE_DATA\]')
    with open('etc\highchart_render_settings\chart4RenderSettings.js') as f:
        contents = f.read()
    contents = cohort_avarage_data.sub(r'78', contents)
    contents = learner_score_data.sub(r'84', contents)
    with open(f'etc\highchart_render_settings\{LEARNER}_chart4RenderSettings.js', 'w') as f:
        f.write(contents)
        
    
    chart1 = os.system(f'.\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_1_Settings.json --outfile .\etc\\temp\chart_images\\new_chart1.png')
    print(chart1)
    # chart2 = os.system(f'.\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_2_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_chart2.png --callback .\etc\\abhijeet.rai\grouped-categories.js')
    # print(chart2)
    # chart3 = os.system(f'.\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_3_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_chart3.png --callback .\etc\\abhijeet.rai\grouped-categories.js')
    # print(chart3)
    # chart4 = os.system(f'.\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_4_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_chart4.png --callback .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_chart4RenderSettings.js')
    # print(chart4)
        
        
    # chart1 = os.system(f'cmd /c c: && .\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_1_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_new_chart1.png')
    # print(chart1)
    # chart2 = os.system(f'cmd /c c: && .\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_2_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_chart2.png --callback .\etc\\abhijeet.rai\grouped-categories.js')
    # print(chart2)
    # chart3 = os.system(f'cmd /c c: && .\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_3_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_chart3.png --callback .\etc\\abhijeet.rai\grouped-categories.js')
    # print(chart3)
    # chart4 = os.system(f'cmd /c c: && .\highcharts-export-server --infile .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_Graph_4_Settings.json --outfile .\etc\\temp\chart_images\{LEARNER}_chart4.png --callback .\etc\\abhijeet.rai\\abhijeet.rai@hcl.com_chart4RenderSettings.js')
    # print(chart4)
        
    # # print(chart, type(chart))       # 4058 file not found         # type <int>
    #                                   # 0 OK
    #                                   # 1 System cant find the path
    #                                   # if file is invalid the server will never stop, no error
    
    # document = Document('media\\docx\\Personal Report Template - NEW FORMAT - 10M.docx')
    
    
    # # FIRST PAGE HEADER
    # trainee_name = document.sections[0].first_page_header.tables[0].rows[1].cells[0].paragraphs[0].runs[0]      # TRAINEE
    # product_name = document.sections[0].first_page_header.tables[0].rows[1].cells[1].paragraphs[0].runs[0]      # Global Business Skills
    # client_cohort = document.sections[0].first_page_header.tables[0].rows[2].cells[0].paragraphs[0].runs[0]     # CLIENT COHORT
    # date = document.sections[0].first_page_header.tables[0].rows[2].cells[1].paragraphs[0].runs[1]              # DATE
        
    # trainee_name.text = trainee_name.text.replace("TRAINEE", "Diana Matkava") 
    # product_name.text = product_name.text.replace("Global Business Skills", "GBS1")
    # client_cohort.text = client_cohort.text.replace("CLIENT COHORT", "ZZL_Cohort 4") 
    # date.text = date.text.replace("DATE", "16.06.2022")
    
    
    # # # PAGE 1 TABLE
    # score_value = document.tables[0].columns[0].cells[1].paragraphs[0].runs[0]   
    # certificat = document.tables[0].columns[1].cells[1].paragraphs[0].runs[0]
    # description1 = document.tables[0].columns[1].cells[3].paragraphs[0].runs[0]
    # description2 = document.tables[0].columns[1].cells[3].paragraphs[0].runs[2]
    # score_value.text = score_value.text.replace("[SCORE]", "100%")
    # certificat.text = certificat.text.replace("[CERT]", "EXELLENT")
    # description1.text = description1.text.replace('[DESCRIPTION1]', 'My custom description 1')
    # description2.text = description2.text.replace('[DESCRIPTION2]', 'My custom description 2')
    
    # # find and replace images in table
    # document.tables[0].columns[0].cells[2]._element.clear_content()
    # document.tables[0].columns[0].cells[2].add_paragraph().add_run().add_picture('media\img\chart4.png', width=Inches(2.5), height=Inches(0.8))
    
    # document.tables[0].columns[1].cells[2]._element.clear_content()
    # document.tables[0].columns[1].cells[2].add_paragraph().add_run().add_picture('media\img\certificate_criteria_distinction.png', width=Inches(1.4), height=Inches(0.5))
    # document.tables[0].columns[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    
 
 
    # # # PARAGRAPHS
    
    # # # Multiple replace in loop
    # context = {
    #     '$cert': 'Excellent',
    #     '$score_heights': 'higher', 
    #     '$score': '100%',
    #     '$user_top_module_1': 'Self Awareness',
    #     '$user_top_module_2': 'Planning and Agility',
    #     '$user_low_module_1': 'Something',
    #     '$user_low_module_2': 'Something',
    # }
    # # document.paragraphs[1].text = reduce(lambda a, kv: a.replace(*kv), context.items(), document.paragraphs[1].text)
   
    # for i in document.paragraphs:
    #     i.text = reduce(lambda a, kv: a.replace(*kv), context.items(), i.text)
        
    
    
    # # works nice if we know indexes
    # # context = {
    # #     'cert': 'Excellent',
    # #     'score_heights': 'higher', 
    # #     'score': 100,
    # #     'user_top_module_1': 'Self Awareness',
    # #     'user_top_module_2': 'Planning and Agility',
    # #     'user_low_module_1': 'Something',
    # #     'user_low_module_2': 'Something',
    # # }
    
    # # text = render_to_string('content.txt', context)
    # # document.paragraphs[1].text = text
    
    
    
    # document.save('media\\docx\\example.docx')
    return HttpResponse()





def report(request):
    
# without charts
#     url = 'http://export.highcharts.com/'
#     headers = {
#     "Accept": "application/json",
#     # "Content-Type": "multipart/form-data",
#   }
    
#     chart_api = {
#         "options": {
#             "chart": {
#                     "type": "bar"
#                 },
#                 "title": {
#                     "text": "Fruit Consumption"
#                 },
#                 "xAxis": {
#                     "categories": ["Apples", "Bananas", "Oranges"]
#                 },
#                 "yAxis": {
#                     "title": {
#                         "text": "Fruit eaten"
#                     }
#                 },
#                 "series": [{
#                     "name": "Jane",
#                     "data": [1, 0, 4]
#                 }, {
#                     "name": "John",
#                     "data": [5, 7, 3]
#                 }]
#             },
        
#         "filename": "test.png",
#         "type": "image/png"
#     #   'async': True
#     }
    
#     r = requests.post(url, data=chart_api, headers=headers)

#     file = open("sample_image.png", "wb")
#     file.write(r.content)
#     file.close()
    
    
    
    # YES Works but has ADs
    doc = aw.Document('csv\\new-Highcharts.html')
    doc.save("html-to-word-output.docx")
    
    # doc = aw.Document()
    # builder = aw.DocumentBuilder(doc)
    # r = requests.get('http://127.0.0.1:4200/new-java-app', allow_redirects=True)
    # print(r.text)
    
    # # urllib.request.urlretrieve("http://127.0.0.1:4200/new-java-app", "test.html")

                
    # builder.insert_html(r.text)
    # doc.save("html-to-word-req.docx")
    
    
    # import urllib.request, urllib.error, urllib.parse

    # url = 'http://127.0.0.1:4200/new-java-app'

    # response = urllib.request.urlopen(url)
    # webContent = response.read().decode('UTF-8')
    # print(webContent[0:300])
    
    
    
    
    
    # from selenium import webdriver
    # from webdriver_manager.chrome import ChromeDriverManager

    # #set chromedriver.exe path
    
    # option = webdriver. ChromeOptions()
    # option. add_argument('headless')
    # driver = webdriver.Chrome(ChromeDriverManager().install(), options=option)
    
    # driver.get("http://127.0.0.1:4200/new-java-app")
    # driver.implicitly_wait(3.5)
    
    # builder.insert_html(driver.page_source)
    # doc.save("html-to-word-sel.docx")
    
    # driver.quit()
    
    
    
    
    # from pywebcopy import save_webpage
    # url = 'http://127.0.0.1:4200/pdf'
    # download_folder = '.\\' 

    # kwargs = {'bypass_robots': True, 'project_name': 'recognisable-name'}

    # save_webpage(url, download_folder, **kwargs)
    
    
    #YES creates file without charts
    # word = win32com.client.Dispatch("Word.Application")
    # in_file  = os.path.abspath(r'csv\\Highcharts.html')
    
    # in_file = requests.get('http://127.0.0.1:4200/new-java-app')
    # in_name  = in_file.text
    # out_file = os.path.abspath("%s123.doc" % in_name)
    
    # doc = word.Documents.Add(in_file)
    # word.Selection.WholeStory()
    # word.Selection.Copy()
    # doc.Close()
    
    # doc = word.Documents.Add()
    # word.Selection.Paste()
    # doc.SaveAs(out_file, FileFormat=0)
    # doc.Close()

    # word.Quit()
    
    
    
    #NO decode error
    # new_parser =
    # HtmlToDocx()
    # new_parser.parse_html_file('csv\Highcharts.html', 'Highcharts-out')
    
    
    
    #NO VEERY BED creates file without charts
    # import pypandoc
    # pypandoc.download_pandoc()
    # output = pypandoc.convert_file('csv\Highcharts.html', 'docx', outputfile="somefile.docx", extra_args=['-RTS'])
    # assert output == ""
    
    
    return HttpResponse()